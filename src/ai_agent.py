import time
import uuid
from datetime import datetime
from typing import Optional, List
from pathlib import Path
import os

from config import Config
from models import ProcessingResult, ProcessingStatus
from audio_processor import AudioProcessor
from llm_processor import LLMProcessor
from calendar_manager import CalendarManager

class AIAgent:
    """Main AI Agent that orchestrates meeting processing and calendar management"""
    
    def __init__(self):
        """Initialize the AI Agent with all components"""
        print("Initializing AI Agent...")
        
        # Validate configuration
        if not Config.validate_config():
            print("Warning: Some configuration is missing. Some features may not work.")
        
        # Initialize components
        self.audio_processor = AudioProcessor()
        self.llm_processor = LLMProcessor()
        self.calendar_manager = CalendarManager()
        
        print("AI Agent initialized successfully")
    
    def process_meeting(self, audio_path: str, meeting_id: Optional[str] = None) -> ProcessingResult:
        """
        Complete meeting processing pipeline
        
        Args:
            audio_path: Path to the audio file
            meeting_id: Optional meeting ID (generated if not provided)
            
        Returns:
            ProcessingResult with all outputs
        """
        start_time = time.time()
        
        if not meeting_id:
            meeting_id = str(uuid.uuid4())
        
        result = ProcessingResult(
            status=ProcessingStatus.PROCESSING,
            meeting_id=meeting_id
        )
        
        try:
            print(f"Starting meeting processing for: {meeting_id}")
            
            # Step 1: Process audio and generate transcript
            print("Step 1: Processing audio and generating transcript...")
            transcript = self.audio_processor.process_audio_file(audio_path, meeting_id)
            result.transcript = transcript
            
            # Step 1.5: Infer speaker names and relabel transcript
            print("Step 1.5: Inferring speaker names and relabeling transcript...")
            speaker_map = self.llm_processor.infer_speaker_names(transcript)
            if speaker_map:
                try:
                    # Relabel segments
                    for segment in transcript.segments:
                        if segment.speaker in speaker_map:
                            segment.speaker = speaker_map[segment.speaker]
                    # Relabel participants
                    for speaker in transcript.speakers:
                        if speaker.speaker_id in speaker_map:
                            speaker.speaker_id = speaker_map[speaker.speaker_id]
                    print(f"✅ Applied speaker names: {speaker_map}")
                except Exception as e:
                    import traceback; traceback.print_exc()
                    raise
            else:
                print("ℹ️  No speaker names could be inferred, keeping original labels")

            # Step 2: Generate meeting minutes using LLM (with speaker names)
            print("Step 2: Generating meeting minutes...")
            minutes = self.llm_processor.generate_meeting_minutes(transcript)
            result.minutes = minutes
            
            # Step 3: Extract calendar events (with speaker names)
            print("Step 3: Extracting calendar events...")
            calendar_events = self.llm_processor.extract_calendar_events(transcript, minutes)
            result.calendar_events = calendar_events
            
            # Step 4: Create calendar events
            if calendar_events:
                print("Step 4: Creating calendar events...")
                created_events = self.calendar_manager.create_events(calendar_events)
                print(f"Created {len(created_events)} calendar events")
            
            # Update result
            result.status = ProcessingStatus.COMPLETED
            result.processing_time = time.time() - start_time
            
            print(f"Meeting processing completed successfully in {result.processing_time:.2f} seconds")
            
        except Exception as e:
            result.status = ProcessingStatus.FAILED
            result.error_message = str(e)
            result.processing_time = time.time() - start_time
            
            print(f"Meeting processing failed: {str(e)}")
        
        return result
    
    def generate_minutes_only(self, audio_path: str, meeting_id: Optional[str] = None) -> ProcessingResult:
        """
        Generate only meeting minutes without calendar events
        
        Args:
            audio_path: Path to the audio file
            meeting_id: Optional meeting ID
            
        Returns:
            ProcessingResult with transcript and minutes
        """
        start_time = time.time()
        
        if not meeting_id:
            meeting_id = str(uuid.uuid4())
        
        result = ProcessingResult(
            status=ProcessingStatus.PROCESSING,
            meeting_id=meeting_id
        )
        
        try:
            print(f"Generating minutes for meeting: {meeting_id}")
            
            # Process audio and generate transcript
            transcript = self.audio_processor.process_audio_file(audio_path, meeting_id)
            result.transcript = transcript
            
            # Generate meeting minutes
            minutes = self.llm_processor.generate_meeting_minutes(transcript)
            result.minutes = minutes
            
            result.status = ProcessingStatus.COMPLETED
            result.processing_time = time.time() - start_time
            
            print(f"Minutes generation completed successfully")
            
        except Exception as e:
            result.status = ProcessingStatus.FAILED
            result.error_message = str(e)
            result.processing_time = time.time() - start_time
            
            print(f"Minutes generation failed: {str(e)}")
        
        return result
    
    def extract_events_only(self, audio_path: str, meeting_id: Optional[str] = None) -> ProcessingResult:
        """
        Extract only calendar events without generating minutes
        
        Args:
            audio_path: Path to the audio file
            meeting_id: Optional meeting ID
            
        Returns:
            ProcessingResult with transcript and calendar events
        """
        start_time = time.time()
        
        if not meeting_id:
            meeting_id = str(uuid.uuid4())
        
        result = ProcessingResult(
            status=ProcessingStatus.PROCESSING,
            meeting_id=meeting_id
        )
        
        try:
            print(f"Extracting events for meeting: {meeting_id}")
            
            # Process audio and generate transcript
            transcript = self.audio_processor.process_audio_file(audio_path, meeting_id)
            result.transcript = transcript
            
            # Extract calendar events
            calendar_events = self.llm_processor.extract_calendar_events(transcript, None)
            result.calendar_events = calendar_events
            
            # Create calendar events
            if calendar_events:
                created_events = self.calendar_manager.create_events(calendar_events)
                print(f"Created {len(created_events)} calendar events")
            
            result.status = ProcessingStatus.COMPLETED
            result.processing_time = time.time() - start_time
            
            print(f"Event extraction completed successfully")
            
        except Exception as e:
            result.status = ProcessingStatus.FAILED
            result.error_message = str(e)
            result.processing_time = time.time() - start_time
            
            print(f"Event extraction failed: {str(e)}")
        
        return result
    
    def generate_transcript_only(self, audio_path: str, meeting_id: Optional[str] = None) -> ProcessingResult:
        """
        Generate only transcript without minutes or calendar events
        
        Args:
            audio_path: Path to the audio file
            meeting_id: Optional meeting ID
            
        Returns:
            ProcessingResult with transcript only
        """
        start_time = time.time()
        
        if not meeting_id:
            meeting_id = str(uuid.uuid4())
        
        result = ProcessingResult(
            status=ProcessingStatus.PROCESSING,
            meeting_id=meeting_id
        )
        
        try:
            print(f"Generating transcript for meeting: {meeting_id}")
            
            # Process audio and generate transcript
            transcript = self.audio_processor.process_audio_file(audio_path, meeting_id)
            result.transcript = transcript
            
            result.status = ProcessingStatus.COMPLETED
            result.processing_time = time.time() - start_time
            
            print(f"Transcript generation completed successfully")
            
        except Exception as e:
            result.status = ProcessingStatus.FAILED
            result.error_message = str(e)
            result.processing_time = time.time() - start_time
            
            print(f"Transcript generation failed: {str(e)}")
        
        return result
    
    def analyze_transcript_file(self, transcript_file: str) -> dict:
        """
        Analyze a transcript file to infer speaker names
        
        Args:
            transcript_file: Path to the transcript file
            
        Returns:
            Dictionary mapping speaker labels to inferred names
        """
        try:
            print(f"Analyzing transcript file: {transcript_file}")
            
            # Analyze the transcript file
            speaker_map = self.analyze_speaker_names_from_file(transcript_file)
            
            return speaker_map
            
        except Exception as e:
            print(f"Error analyzing transcript file: {str(e)}")
            raise
    
    def save_transcript_to_file(self, transcript, output_path: str = ""):
        """
        Save transcript to a text file in the transcripts directory
        
        Args:
            transcript: MeetingTranscript object
            output_path: Optional custom path (if empty, uses default transcripts directory)
        """
        try:
            # Use default transcripts directory if no path specified
            if not output_path:
                output_path = os.path.join(Config.TRANSCRIPTS_DIR, f"{transcript.meeting_id}_transcript.txt")
            
            with open(output_path, 'w', encoding='utf-8') as f:
                # Write header
                f.write(f"Meeting Transcript - {transcript.meeting_id}\n")
                f.write(f"Date: {transcript.created_at.strftime('%Y-%m-%d %H:%M:%S')}\n")
                f.write(f"Duration: {transcript.duration//1000//60:02d}:{transcript.duration//1000%60:02d}\n")
                f.write(f"Participants: {', '.join([f'{s.speaker_id}' for s in transcript.speakers])}\n")
                f.write("=" * 80 + "\n\n")
                
                # Write transcript segments
                for segment in transcript.segments:
                    timestamp = f"[{segment.start//1000//60:02d}:{segment.start//1000%60:02d}]"
                    f.write(f"{timestamp} {segment.speaker}: {segment.text}\n")
                
            print(f"✅ Transcript saved to: {output_path}")
            
        except Exception as e:
            print(f"❌ Error saving transcript: {str(e)}")
            raise

    def save_minutes_to_file(self, minutes, output_path: str = ""):
        """
        Save meeting minutes to a DOCX file in the minutes directory
        
        Args:
            minutes: MeetingMinutes object
            output_path: Optional custom path (if empty, uses default minutes directory)
        """
        try:
            # Use default minutes directory if no path specified
            if not output_path:
                output_path = os.path.join(Config.MINUTES_DIR, f"{minutes.meeting_id}_minutes.docx")
            
            # Always save as DOCX
            self.save_minutes_to_docx(minutes, output_path)
            
        except Exception as e:
            print(f"❌ Error saving meeting minutes: {str(e)}")
            raise
    
    def save_minutes_to_docx(self, minutes, output_path: str):
        """
        Save meeting minutes to a well-structured and formatted DOCX file
        Args:
            minutes: MeetingMinutes object
            output_path: Path to save the docx file
        """
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.shared import RGBColor

        doc = Document()
        # Title
        title = doc.add_heading('Meeting Minutes', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Date and Duration
        doc.add_paragraph(f"Date: {minutes.date.strftime('%A, %B %d, %Y %I:%M %p')}")
        doc.add_paragraph(f"Duration: {minutes.duration.total_seconds() / 60:.1f} minutes")
        doc.add_paragraph(f"Participants: {', '.join(minutes.participants)}")
        doc.add_paragraph("")

        # Key Points
        doc.add_heading('Key Points Discussed', level=1)
        if minutes.key_points:
            for point in minutes.key_points:
                doc.add_paragraph(point, style='List Bullet')
        else:
            doc.add_paragraph("None.")

        # Action Items
        doc.add_heading('Action Items', level=1)
        if minutes.action_items:
            for item in minutes.action_items:
                p = doc.add_paragraph(style='List Number')
                p.add_run(f"{item.description}").bold = True
                p.add_run(f" (Assignee: {item.assignee or 'TBD'}, Due: {item.due_date.strftime('%Y-%m-%d') if item.due_date else 'TBD'}, Priority: {item.priority}, Status: {item.status.capitalize()})")
        else:
            doc.add_paragraph("None.")

        # Decisions
        doc.add_heading('Decisions Made', level=1)
        if minutes.decisions:
            for decision in minutes.decisions:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(f"{decision.topic}: ").bold = True
                p.add_run(f"{decision.decision}")
                if decision.rationale:
                    p.add_run(f" (Rationale: {decision.rationale})").italic = True
        else:
            doc.add_paragraph("None.")

        # Next Steps
        doc.add_heading('Next Steps', level=1)
        if minutes.next_steps:
            for step in minutes.next_steps:
                doc.add_paragraph(step, style='List Bullet')
        else:
            doc.add_paragraph("None.")

        # Summary
        doc.add_heading('Summary', level=1)
        doc.add_paragraph(minutes.summary or "None.")

        doc.save(output_path)
        print(f"Meeting minutes saved to: {output_path}")
    
    def get_processing_summary(self, result: ProcessingResult) -> str:
        """
        Get a summary of the processing results
        
        Args:
            result: ProcessingResult object
            
        Returns:
            Formatted summary string
        """
        summary = f"""
=== Meeting Processing Summary ===
Meeting ID: {result.meeting_id}
Status: {result.status.value}
Processing Time: {result.processing_time:.2f} seconds

"""
        
        if result.error_message:
            summary += f"Error: {result.error_message}\n"
        
        if result.transcript:
            summary += f"""
Transcript:
- Duration: {result.transcript.duration / 1000 / 60:.1f} minutes
- Speakers: {len(result.transcript.speakers)}
- Segments: {len(result.transcript.segments)}
"""
        
        if result.minutes:
            summary += f"""
Meeting Minutes:
- Key Points: {len(result.minutes.key_points)}
- Action Items: {len(result.minutes.action_items)}
- Decisions: {len(result.minutes.decisions)}
- Next Steps: {len(result.minutes.next_steps)}
"""
        
        if result.calendar_events:
            summary += f"""
Calendar Events:
- Events Created: {len(result.calendar_events)}
"""
        
        return summary 

    def analyze_speaker_names_from_file(self, transcript_file: str) -> dict:
        """
        Analyze a saved transcript file to infer speaker names
        
        Args:
            transcript_file: Path to the transcript file
            
        Returns:
            Dictionary mapping speaker labels to inferred names
        """
        try:
            # Read the transcript file
            with open(transcript_file, 'r', encoding='utf-8') as f:
                transcript_content = f.read()
            
            print(f"📖 Analyzing transcript file: {transcript_file}")
            
            # Use LLM to analyze the transcript and infer speaker names
            speaker_map = self.llm_processor.analyze_speaker_names_from_text(transcript_content)
            
            print("🔍 Speaker name analysis completed")
            print("📊 Inferred speaker mapping:")
            for speaker_label, inferred_name in speaker_map.items():
                print(f"   {speaker_label} → {inferred_name}")
            
            return speaker_map
            
        except Exception as e:
            print(f"❌ Error analyzing transcript file: {str(e)}")
            raise 
